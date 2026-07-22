from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from .models import ReportSummary


class FillReportWriter:
    def write(
        self,
        summary: ReportSummary,
        docx_path: str | Path,
        json_path: str | Path | None = None,
    ) -> tuple[Path, Path | None]:
        docx_path = Path(docx_path)
        docx_path.parent.mkdir(parents=True, exist_ok=True)
        document = Document()
        title = document.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("填充报告")
        run.bold = True
        run.font.size = Pt(18)
        document.add_paragraph(
            f"字段总数：{summary.total_fields}；已填：{summary.filled_fields}；"
            f"为空：{summary.empty_fields}；失败：{summary.failed_fields}；"
            f"格式异常：{summary.abnormal_fields}；未知输入：{summary.unknown_input_fields}"
        )
        pv = summary.paragraph_validation
        document.add_paragraph(
            f"段落数验证：填充前 {pv.before}，填充后 {pv.after}，"
            f"预期增量 {pv.expected_delta}，实际增量 {pv.actual_delta}，"
            f"允许偏差 ±{pv.tolerance}，结果：{'通过' if pv.success else '失败'}。"
        )
        table = document.add_table(rows=1, cols=7)
        table.style = "Table Grid"
        headers = ["状态", "字段", "类别", "填充值", "位置", "格式异常", "说明"]
        for cell, value in zip(table.rows[0].cells, headers):
            cell.text = value
        for item in summary.results:
            row = table.add_row().cells
            status = "已填" if item.success and item.filled != "" else ("为空" if item.success else "失败")
            values = [
                status,
                item.field,
                item.category,
                item.filled,
                item.location,
                "是" if item.format_abnormal else "否",
                item.message,
            ]
            for cell, value in zip(row, values):
                cell.text = value
        document.save(docx_path)
        written_json = None
        if json_path is not None:
            written_json = Path(json_path)
            written_json.parent.mkdir(parents=True, exist_ok=True)
            written_json.write_text(
                json.dumps(summary.to_dict(), ensure_ascii=False, indent=2),
                encoding="utf-8",
            )
        return docx_path, written_json
