from __future__ import annotations

import json
import sys
import tempfile
import unittest
import zipfile
from pathlib import Path

from docx import Document
from lxml import etree
from PIL import Image

PACKAGE_ROOT = Path(__file__).resolve().parents[2]
sys.path.insert(0, str(PACKAGE_ROOT))

from docx_engine import DocxGenerator, StyleSpec
from docx_engine.checkbox_filler import CheckboxFiller
from docx_engine.cover_filler import CoverFiller
from docx_engine.package import DocxPackage
from docx_engine.qa_filler import QASectionFiller
from docx_engine.style_manager import StyleManager
from docx_engine.table_filler import TableFiller


class EngineTestCase(unittest.TestCase):
    @classmethod
    def setUpClass(cls):
        local_template = PACKAGE_ROOT / "FOF尽调报告_书签模板.docx"
        cls.template = (
            local_template
            if local_template.exists()
            else Path("/Users/a/Desktop/任务/第二阶段/1-1/2.2+2.3/FOF尽调报告_书签模板.docx")
        )
        cls.manifest = PACKAGE_ROOT / "bookmark_manifest.json"
        assert cls.template.exists()
        assert cls.manifest.exists()

    def test_end_to_end_and_report(self):
        with tempfile.TemporaryDirectory() as td:
            output = Path(td) / "result.docx"
            generator = DocxGenerator(self.template, self.manifest)
            result = generator.generate(
                {
                    "cover_manager_name": "测试管理人",
                    "cover_investigator": "测试员",
                    "cover_strategy_futures_quant_trend": True,
                    "table_1_row0_col1": "测试管理人",
                    "qa_section1_q001_answer": "回答第一段\n回答第二段",
                },
                output,
            )
            self.assertTrue(output.exists())
            self.assertTrue(result.report_docx.exists())
            self.assertTrue(result.report_json.exists())
            self.assertTrue(result.summary.paragraph_validation.success)
            report = json.loads(result.report_json.read_text(encoding="utf-8"))
            self.assertGreater(report["filled_fields"], 0)
            text = "\n".join(p.text for p in Document(output).paragraphs)
            self.assertIn("测试管理人", text)
            self.assertIn("☑期货量化趋势", text)
            self.assertIn("回答第二段", text)
            self.assertIn("调查人员：测试员", text)
            self.assertNotIn("调查人员：____________________", text)
            self.assertEqual(
                result.summary.paragraph_validation.actual_delta,
                result.summary.paragraph_validation.expected_delta,
            )

    def test_operation_tuple_contract(self):
        package = DocxPackage(self.template)
        try:
            bookmarks = package.bookmark_index()
            styles = StyleManager(package)
            operations = [
                CoverFiller(package, bookmarks, styles).fill("cover_manager_name", "测试"),
                CheckboxFiller(package, bookmarks, styles).fill("cover_strategy_composite", True),
                TableFiller(package, bookmarks, styles).fill(1, 0, 1, "测试"),
                QASectionFiller(package, bookmarks, styles).fill("qa_section1_q001_answer", "测试"),
                TableFiller(package, bookmarks, styles).merge(2, (5, 0), (5, 1)),
            ]
            for result in operations:
                self.assertEqual(len(result), 3)
                self.assertTrue(result[0])
        finally:
            package.close()

    def test_unknown_field_is_reported(self):
        with tempfile.TemporaryDirectory() as td:
            result = DocxGenerator(self.template, self.manifest).generate(
                {"not_a_bookmark": "x"},
                Path(td) / "unknown.docx",
            )
            self.assertEqual(result.summary.unknown_input_fields, 1)
            self.assertGreaterEqual(result.summary.failed_fields, 1)

    def test_format_abnormal_is_reported(self):
        with tempfile.TemporaryDirectory() as td:
            result = DocxGenerator(
                self.template,
                self.manifest,
                styles={"cover": StyleSpec(color="NOT_RGB")},
            ).generate(
                {"cover_manager_name": "测试"},
                Path(td) / "bad_style.docx",
            )
            self.assertGreaterEqual(result.summary.abnormal_fields, 1)

    def test_table_merge(self):
        with tempfile.TemporaryDirectory() as td:
            output = Path(td) / "merge.docx"
            result = DocxGenerator(self.template, self.manifest).generate(
                {
                    "table_2_row5_col0": "合并内容",
                    "merges": [
                        {"table": 2, "start": [5, 0], "end": [5, 1]}
                    ],
                },
                output,
            )
            self.assertTrue(output.exists())
            self.assertFalse(any(not x.success for x in result.summary.results if x.field == "table_merge"))
            document = Document(output)
            self.assertEqual(len(document.tables[1].rows[5].cells), 4)
            self.assertEqual(document.tables[1].cell(5, 0).text, document.tables[1].cell(5, 1).text)

    def test_table_multiline_value_uses_word_line_break(self):
        with tempfile.TemporaryDirectory() as td:
            output = Path(td) / "multiline.docx"
            DocxGenerator(self.template, self.manifest).generate(
                {"table_1_row3_col1": "注册地址\n办公地址"},
                output,
            )
            document = Document(output)
            self.assertEqual(
                document.tables[0].cell(3, 1).text,
                "注册地址\n办公地址",
            )
            with zipfile.ZipFile(output) as archive:
                root = etree.fromstring(archive.read("word/document.xml"))
            ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
            bookmark = root.xpath(
                "//w:bookmarkStart[@w:name='table_1_row3_col1']",
                namespaces=ns,
            )[0]
            self.assertTrue(
                bookmark.getparent().xpath(".//w:br", namespaces=ns)
            )

    def test_four_strategy_routes(self):
        cases = {
            "quant": ("factor_entry", "量化专属回答"),
            "cta": ("factor_entry", "CTA专属回答"),
            "bond": ("investment_scope", "债券专属回答"),
            "option": ("volatility_forecast", "期权专属回答"),
        }
        with tempfile.TemporaryDirectory() as td:
            for selected, (field, marker) in cases.items():
                strategies = {
                    name: {"_applicable": name == selected}
                    for name in cases
                }
                strategies[selected][field] = marker
                output = Path(td) / f"{selected}.docx"
                result = DocxGenerator(self.template, self.manifest).generate(
                    {"strategies": strategies},
                    output,
                )
                text = "\n".join(p.text for p in Document(output).paragraphs)
                self.assertIn(marker, text)
                for other, (_, other_marker) in cases.items():
                    if other != selected:
                        self.assertNotIn(other_marker, text)
                self.assertTrue(result.summary.paragraph_validation.success)
                skipped = [
                    item for item in result.summary.results
                    if item.field.startswith("strat_")
                    and item.category == "qa"
                    and "策略不适用" in item.message
                ]
                self.assertGreater(len(skipped), 0)

    def test_answer_after_numbered_question_is_not_numbered(self):
        with tempfile.TemporaryDirectory() as td:
            output = Path(td) / "numbering.docx"
            marker = "主观多头策略回答不应带序号"
            blank_marker = "空白回答段落也不应带序号"
            hanging_marker = "CTA回答不得继承字符悬挂缩进"
            DocxGenerator(self.template, self.manifest).generate(
                {
                    "qa_section2_q024_answer": marker,
                    "qa_section1_q002_answer": blank_marker,
                    "strat_cta_position_sizing": hanging_marker,
                },
                output,
            )
            with zipfile.ZipFile(output) as archive:
                root = etree.fromstring(archive.read("word/document.xml"))
            ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
            paragraphs = root.xpath("//w:body//w:p", namespaces=ns)
            for expected in (marker, blank_marker, hanging_marker):
                answer = next(
                    paragraph
                    for paragraph in paragraphs
                    if expected in "".join(paragraph.xpath(".//w:t/text()", namespaces=ns))
                )
                self.assertFalse(answer.xpath("./w:pPr/w:numPr", namespaces=ns))
                indentation = answer.find(
                    "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr/"
                    "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ind"
                )
                self.assertIsNotNone(indentation)
                word = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
                self.assertEqual(indentation.get(word + "leftChars"), "0")
                self.assertEqual(indentation.get(word + "firstLineChars"), "200")
                self.assertIsNone(indentation.get(word + "hanging"))
                self.assertIsNone(indentation.get(word + "hangingChars"))

    def test_risk_answer_is_inserted_after_all_listed_scenarios(self):
        with tempfile.TemporaryDirectory() as td:
            output = Path(td) / "risk_answer.docx"
            marker = "无"
            DocxGenerator(self.template, self.manifest).generate(
                {"qa_section3_q122_answer": marker},
                output,
            )
            document = Document(output)
            texts = [paragraph.text.strip() for paragraph in document.paragraphs]
            final_scenario = next(
                index
                for index, text in enumerate(texts)
                if text.startswith("（5）有证据表明管理人未来持续经营存在重大不确定性")
            )
            answer = next(
                index
                for index in range(final_scenario + 1, len(texts))
                if texts[index] == marker
            )
            next_section = next(
                index
                for index, text in enumerate(texts)
                if index > final_scenario and text == "产品管理情况"
            )
            self.assertGreater(answer, final_scenario)
            self.assertLess(answer, next_section)

    def test_revised_template_inserts_bookmarked_image(self):
        revised_template = PACKAGE_ROOT / "FOF尽调报告_书签模板_2026修订版.docx"
        revised_manifest = PACKAGE_ROOT / "bookmark_manifest_2026修订版.json"
        if not revised_template.exists() or not revised_manifest.exists():
            self.skipTest("revised template is not available")
        with tempfile.TemporaryDirectory() as td:
            image = Path(td) / "org.png"
            Image.new("RGB", (600, 300), "white").save(image)
            output = Path(td) / "image.docx"
            result = DocxGenerator(revised_template, revised_manifest).generate(
                {"image_org_structure": str(image)},
                output,
            )
            self.assertEqual(result.summary.failed_fields, 0)
            with zipfile.ZipFile(output) as archive:
                root = etree.fromstring(archive.read("word/document.xml"))
            ns = {
                "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
            }
            bookmark = root.xpath(
                "//w:bookmarkStart[@w:name='image_org_structure']",
                namespaces=ns,
            )[0]
            self.assertTrue(
                bookmark.getparent().xpath(".//w:drawing", namespaces=ns)
            )


if __name__ == "__main__":
    unittest.main()
