from __future__ import annotations

import sys
import tempfile
import unittest
from pathlib import Path

from docx import Document

OUTPUTS = Path(__file__).resolve().parents[2]
sys.path.insert(0, str(OUTPUTS))

from licensed_bookmark_template_engine import BookmarkTemplateEngine


class LicensedTemplateTestCase(unittest.TestCase):
    @classmethod
    def setUpClass(cls):
        cls.template = OUTPUTS / "FOF尽调报告_持牌金融机构_书签模板.docx"
        cls.manifest = OUTPUTS / "bookmark_manifest_持牌金融机构.json"
        cls.data = OUTPUTS / "example_data_持牌金融机构.json"

    def test_bookmark_inventory(self):
        engine = BookmarkTemplateEngine(self.template, self.manifest)
        names = engine.list_bookmarks()
        self.assertEqual(len(names), 69)
        self.assertEqual(len(names), len(set(names)))
        self.assertIn("cover_manager_name", names)
        self.assertIn("table_1_row0_col1", names)
        self.assertIn("qa_section5_conclusion", names)

    def test_json_render(self):
        with tempfile.TemporaryDirectory() as td:
            output = Path(td) / "licensed.docx"
            result = BookmarkTemplateEngine(
                self.template,
                self.manifest,
            ).render_json(self.data, output)
            self.assertTrue(output.exists())
            self.assertTrue(result.report_docx.exists())
            self.assertTrue(result.report_json.exists())
            self.assertTrue(result.summary.paragraph_validation.success)
            document = Document(output)
            body_text = "\n".join(p.text for p in document.paragraphs)
            self.assertIn("示例持牌金融机构", body_text)
            self.assertIn("☑股票量化选股", body_text)
            self.assertIn("采用多因子量化选股策略。", body_text)
            self.assertEqual(document.tables[0].cell(0, 1).text, "示例持牌金融机构")


if __name__ == "__main__":
    unittest.main()
