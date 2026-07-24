from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Emu, Pt, RGBColor


NAVY = "1F4E78"
LIGHT_BLUE = "D9EAF7"
LIGHT_GRAY = "F2F4F7"


def _set_cell_fill(cell, color: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), color)


def _set_cell_width(cell, width_dxa: int) -> None:
    properties = cell._tc.get_or_add_tcPr()
    width = properties.find(qn("w:tcW"))
    if width is None:
        width = OxmlElement("w:tcW")
        properties.append(width)
    width.set(qn("w:type"), "dxa")
    width.set(qn("w:w"), str(width_dxa))


def _set_table_geometry(table, widths_dxa: list[int]) -> None:
    table.autofit = False
    properties = table._tbl.tblPr
    table_width = properties.find(qn("w:tblW"))
    if table_width is None:
        table_width = OxmlElement("w:tblW")
        properties.append(table_width)
    table_width.set(qn("w:type"), "dxa")
    table_width.set(qn("w:w"), str(sum(widths_dxa)))

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width_dxa in widths_dxa:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width_dxa))
        grid.append(column)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            _set_cell_width(cell, widths_dxa[index])
            cell.width = Emu(widths_dxa[index] * 635)


def _repeat_header(row) -> None:
    properties = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    properties.append(repeat)


def _set_cell_text(
    cell,
    value: Any,
    *,
    bold: bool = False,
    color: RGBColor | None = None,
    size: float = 9,
    align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT,
) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.05
    run = paragraph.add_run(str(value))
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "SimSun"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "宋体")
    if color is not None:
        run.font.color.rgb = color
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def append_scorecard(path: Path, scorecard: dict[str, Any]) -> None:
    document = Document(path)
    section = document.sections[-1]
    usable_width_dxa = int(
        (section.page_width - section.left_margin - section.right_margin) / 635
    )

    document.add_page_break()
    heading = document.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.paragraph_format.space_after = Pt(10)
    run = heading.add_run("附录：私募产品准入评分表")
    run.bold = True
    run.font.size = Pt(15)
    run.font.color.rgb = RGBColor(31, 78, 120)
    run.font.name = "SimSun"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "宋体")

    metrics = scorecard.get("metrics") or {}
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(8)
    subtitle_run = subtitle.add_run(
        f"计算期间：{metrics.get('start_date', '—')} 至 {metrics.get('end_date', '—')}　"
        f"波动率分类：{metrics.get('volatility_band_label', '—')}"
    )
    subtitle_run.font.size = Pt(9)
    subtitle_run.font.name = "SimSun"
    subtitle_run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "宋体")

    summary = document.add_table(rows=2, cols=4)
    summary.style = "Table Grid"
    summary_widths = [
        usable_width_dxa // 4,
        usable_width_dxa // 4,
        usable_width_dxa // 4,
        usable_width_dxa - 3 * (usable_width_dxa // 4),
    ]
    _set_table_geometry(summary, summary_widths)
    summary_values = [
        ("定量得分", f"{scorecard.get('quantitative_score', 0):g} / 62"),
        ("定性得分", f"{scorecard.get('qualitative_score', 0):g} / 38"),
        ("合规扣分", f"-{scorecard.get('compliance_deduction', 0):g}"),
        ("总分", f"{scorecard.get('total_score', 0):g} / 100"),
    ]
    for index, (label, value) in enumerate(summary_values):
        _set_cell_fill(summary.rows[0].cells[index], NAVY)
        _set_cell_text(
            summary.rows[0].cells[index],
            label,
            bold=True,
            color=RGBColor(255, 255, 255),
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
        _set_cell_fill(summary.rows[1].cells[index], LIGHT_BLUE)
        _set_cell_text(
            summary.rows[1].cells[index],
            value,
            bold=True,
            size=10,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )

    decision = document.add_paragraph()
    decision.alignment = WD_ALIGN_PARAGRAPH.CENTER
    decision.paragraph_format.space_before = Pt(7)
    decision.paragraph_format.space_after = Pt(9)
    decision_run = decision.add_run(
        "评分结论："
        + ("达到入池标准（总分≥60分）" if scorecard.get("admitted") else "未达到入池标准（总分<60分）")
    )
    decision_run.bold = True
    decision_run.font.size = Pt(10)
    decision_run.font.color.rgb = (
        RGBColor(20, 115, 70) if scorecard.get("admitted") else RGBColor(180, 45, 45)
    )
    decision_run.font.name = "SimSun"
    decision_run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "宋体")

    rows = list(scorecard.get("score_rows") or [])
    details = document.add_table(rows=1, cols=5)
    details.style = "Table Grid"
    widths = [
        int(usable_width_dxa * 0.13),
        int(usable_width_dxa * 0.20),
        int(usable_width_dxa * 0.20),
        int(usable_width_dxa * 0.10),
    ]
    widths.append(usable_width_dxa - sum(widths))
    _set_table_geometry(details, widths)
    headers = ("一级维度", "二级指标", "指标值", "得分", "评分依据")
    for index, label in enumerate(headers):
        _set_cell_fill(details.rows[0].cells[index], NAVY)
        _set_cell_text(
            details.rows[0].cells[index],
            label,
            bold=True,
            color=RGBColor(255, 255, 255),
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
    _repeat_header(details.rows[0])

    for row_index, item in enumerate(rows):
        cells = details.add_row().cells
        maximum = item.get("maximum", "")
        score_text = (
            f"{item.get('score', 0):g}"
            if maximum == "扣分"
            else f"{item.get('score', 0):g}/{maximum}"
        )
        values = (
            item.get("category", ""),
            item.get("indicator", ""),
            item.get("value", ""),
            score_text,
            item.get("basis", ""),
        )
        for index, value in enumerate(values):
            if row_index % 2:
                _set_cell_fill(cells[index], LIGHT_GRAY)
            _set_cell_text(
                cells[index],
                value,
                align=WD_ALIGN_PARAGRAPH.CENTER if index in {0, 3} else WD_ALIGN_PARAGRAPH.LEFT,
            )
    _set_table_geometry(details, widths)

    source = document.add_paragraph()
    source.paragraph_format.space_before = Pt(7)
    source.paragraph_format.space_after = Pt(0)
    source_run = source.add_run(
        "规则来源：开源证券私募产品准入打分卡（波动率区分版）。"
        "定量指标由上传净值计算；定性指标来自评分卡结构化表单。"
    )
    source_run.font.size = Pt(8)
    source_run.font.color.rgb = RGBColor(95, 105, 120)
    source_run.font.name = "SimSun"
    source_run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "宋体")
    document.save(path)
