from datetime import date
from io import BytesIO
from pathlib import Path
from xml.etree import ElementTree
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document

from app import storage
from app.services.scorecard import parse_nav_upload
from app.services.scorecard_excel import render_scorecard_workbook
from tests.test_api_workflow import create_report_records


def _minimal_nav_xlsx() -> bytes:
    output = BytesIO()
    with ZipFile(output, "w", ZIP_DEFLATED) as archive:
        archive.writestr(
            "[Content_Types].xml",
            """<?xml version="1.0" encoding="UTF-8"?>
            <Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
              <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
              <Default Extension="xml" ContentType="application/xml"/>
              <Override PartName="/xl/workbook.xml" ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet.main+xml"/>
              <Override PartName="/xl/worksheets/sheet1.xml" ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.worksheet+xml"/>
            </Types>""",
        )
        archive.writestr(
            "_rels/.rels",
            """<?xml version="1.0" encoding="UTF-8"?>
            <Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
              <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="xl/workbook.xml"/>
            </Relationships>""",
        )
        archive.writestr(
            "xl/workbook.xml",
            """<?xml version="1.0" encoding="UTF-8"?>
            <workbook xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main"
              xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">
              <sheets><sheet name="净值" sheetId="1" r:id="rId1"/></sheets>
            </workbook>""",
        )
        archive.writestr(
            "xl/_rels/workbook.xml.rels",
            """<?xml version="1.0" encoding="UTF-8"?>
            <Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
              <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/worksheet" Target="worksheets/sheet1.xml"/>
            </Relationships>""",
        )
        archive.writestr(
            "xl/worksheets/sheet1.xml",
            """<?xml version="1.0" encoding="UTF-8"?>
            <worksheet xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main">
              <sheetData>
                <row r="1">
                  <c r="A1" t="inlineStr"><is><t>日期</t></is></c>
                  <c r="B1" t="inlineStr"><is><t>累计净值</t></is></c>
                </row>
                <row r="2">
                  <c r="A2" t="inlineStr"><is><t>2025-01-01</t></is></c>
                  <c r="B2"><v>1.0</v></c>
                </row>
                <row r="3">
                  <c r="A3" t="inlineStr"><is><t>2025-02-01</t></is></c>
                  <c r="B3"><v>1.02</v></c>
                </row>
              </sheetData>
            </worksheet>""",
        )
    return output.getvalue()


def _nav_csv() -> bytes:
    values = [
        1.0000, 1.0200, 1.0404, 1.0612, 1.0506, 1.0716, 1.0930,
        1.1149, 1.1038, 1.1258, 1.1483, 1.1713, 1.1947, 1.2186,
        1.2429, 1.2678, 1.2931, 1.3189, 1.3453, 1.3722,
    ]
    rows = ["日期,累计净值"]
    year, month = 2024, 1
    for value in values:
        rows.append(f"{date(year, month, 1).isoformat()},{value:.4f}")
        month += 1
        if month == 13:
            year += 1
            month = 1
    return ("\n".join(rows) + "\n").encode()


def _calculation_payload() -> dict:
    return {
        "date_column": "日期",
        "nav_column": "累计净值",
        "benchmark_column": None,
        "benchmark_mode": "absolute",
        "risk_free_rate_percent": 0,
        "qualitative": {
            "strategy_scale_group": "cta_t0",
            "managed_scale_100m": 150,
            "active_product_count": 115,
            "company_headcount": 65,
            "manager_same_strategy_years": 5,
            "manager_industry_years": 10,
            "manager_philosophy_level": "complete",
            "manager_profile_stable": True,
            "research_headcount": 48,
            "research_background_match": True,
            "core_research_experience_years": 5,
            "research_live_track_record": True,
            "core_departures_1y": 0,
            "core_departures_3y": 0,
            "incentive_level": "long_term",
            "current_strategy_scale_100m": 55,
            "theoretical_capacity_100m": 90,
            "differentiation_level": "significant",
            "risk_system_level": "complete",
            "risk_team_headcount": 3,
            "risk_team_experience_years": 11,
            "manager_coinvest_percent": 0,
            "manager_coinvest_lock_years": 0,
            "core_personal_coinvest": False,
            "regulatory_events_3y": 0,
            "negative_or_litigation_events_3y": 0,
        },
    }


def test_xlsx_parser_reads_inline_strings_and_numbers() -> None:
    table = parse_nav_upload(_minimal_nav_xlsx(), ".xlsx")
    assert table.sheet_name == "净值"
    assert table.detected_columns == {
        "date": "日期",
        "nav": "累计净值",
        "benchmark": None,
    }
    assert table.rows[1]["累计净值"] == 1.02


def _manual_scores() -> dict[str, float]:
    return {
        "one_year_return": 10, "relative_return": 9, "long_term_return": 7,
        "monthly_win_rate": 6, "max_drawdown": 8, "sharpe_ratio": 9, "calmar_ratio": 6,
        "managed_products": 10, "investment_manager": 5, "research_team": 4,
        "team_stability": 2, "allocation_value": 3, "risk_control": 3,
        "coinvestment": 2, "compliance_deduction": 3,
    }


def _xlsx_numeric_cells(path: Path) -> dict[str, float]:
    namespace = {"x": "http://schemas.openxmlformats.org/spreadsheetml/2006/main"}
    with ZipFile(path) as archive:
        root = ElementTree.fromstring(archive.read("xl/worksheets/sheet1.xml"))
    return {
        cell.attrib["r"]: float(cell.findtext("x:v", namespaces=namespace))
        for cell in root.findall(".//x:c", namespace)
        if cell.attrib.get("r", "").startswith("D") and cell.find("x:v", namespace) is not None
    }


def test_scorecard_api_calculates_and_keeps_word_independent(
    client,
    private_fund_data,
    tmp_path: Path,
    monkeypatch,
) -> None:
    generated_dir = tmp_path / "generated"
    nav_dir = tmp_path / "nav"
    monkeypatch.setattr(storage, "GENERATED_DIR", generated_dir)
    monkeypatch.setattr(storage, "NAV_UPLOAD_DIR", nav_dir)
    _, _, report = create_report_records(
        client,
        private_fund_data,
        "private_fund",
        "评分卡端到端测试",
    )

    uploaded = client.post(
        f"/api/reports/{report['id']}/scorecard/nav",
        content=_nav_csv(),
        headers={"Content-Type": "text/csv", "X-Filename": "nav.csv"},
    )
    assert uploaded.status_code == 201, uploaded.text
    assert uploaded.json()["detected_columns"]["date"] == "日期"
    assert uploaded.json()["detected_columns"]["nav"] == "累计净值"

    calculated = client.post(
        f"/api/reports/{report['id']}/scorecard/calculate",
        json=_calculation_payload(),
    )
    assert calculated.status_code == 200, calculated.text
    result = calculated.json()
    assert result["quantitative_score"] <= 62
    assert result["qualitative_score"] == 33
    assert result["total_score"] == (
        result["quantitative_score"] + result["qualitative_score"]
    )
    assert len(result["score_rows"]) == 14
    assert result["metrics"]["observations"] == 20

    generation = client.post(f"/api/reports/{report['id']}/generate")
    assert generation.status_code == 200, generation.text
    output_path = generated_dir / generation.json()["filename"]
    document = Document(output_path)
    assert not any(paragraph.text == "附录：私募产品准入评分表" for paragraph in document.paragraphs)


def test_licensed_template_word_also_stays_independent_from_scorecard(
    client,
    licensed_institution_data,
    tmp_path: Path,
    monkeypatch,
) -> None:
    generated_dir = tmp_path / "licensed-generated"
    nav_dir = tmp_path / "licensed-nav"
    monkeypatch.setattr(storage, "GENERATED_DIR", generated_dir)
    monkeypatch.setattr(storage, "NAV_UPLOAD_DIR", nav_dir)
    _, _, report = create_report_records(
        client,
        licensed_institution_data,
        "licensed_institution",
        "持牌机构评分卡测试",
    )
    uploaded = client.post(
        f"/api/reports/{report['id']}/scorecard/nav",
        content=_nav_csv(),
        headers={"Content-Type": "text/csv; charset=utf-8", "X-Filename": "nav.csv"},
    )
    assert uploaded.status_code == 201, uploaded.text
    calculated = client.post(
        f"/api/reports/{report['id']}/scorecard/calculate",
        json=_calculation_payload(),
    )
    assert calculated.status_code == 200, calculated.text
    generation = client.post(f"/api/reports/{report['id']}/generate")
    assert generation.status_code == 200, generation.text
    document = Document(generated_dir / generation.json()["filename"])
    assert not any(paragraph.text == "附录：私募产品准入评分表" for paragraph in document.paragraphs)


def test_manual_scores_generate_filled_excel(client, private_fund_data, tmp_path: Path, monkeypatch) -> None:
    scorecard_dir = tmp_path / "scorecards"
    monkeypatch.setattr(storage, "SCORECARD_GENERATED_DIR", scorecard_dir)
    _, _, report = create_report_records(client, private_fund_data, "private_fund", "人工评分Excel测试")
    initial = client.get(f"/api/reports/{report['id']}/scorecard")
    assert initial.status_code == 200
    assert len(initial.json()["template_items"]) == 15
    saved = client.put(f"/api/reports/{report['id']}/scorecard/manual", json={"scores": _manual_scores()})
    assert saved.status_code == 200, saved.text
    assert saved.json()["quantitative_score"] == 46
    assert saved.json()["qualitative_score"] == 29
    assert saved.json()["total_score"] == 72
    generated = client.post(f"/api/reports/{report['id']}/scorecard/generate-excel")
    assert generated.status_code == 200, generated.text
    path = scorecard_dir / generated.json()["filename"]
    cells = _xlsx_numeric_cells(path)
    assert cells["D4"] == 10 and cells["D5"] == 9 and cells["D18"] == 3 and cells["D19"] == 72
    downloaded = client.get(generated.json()["download_url"])
    assert downloaded.status_code == 200


def test_manual_score_validation_rejects_over_maximum(client, private_fund_data) -> None:
    _, _, report = create_report_records(client, private_fund_data, "private_fund", "人工评分上限测试")
    scores = _manual_scores()
    scores["investment_manager"] = 7
    response = client.put(f"/api/reports/{report['id']}/scorecard/manual", json={"scores": scores})
    assert response.status_code == 422
    assert "不能超过满分6分" in response.text


def test_template_renderer_preserves_workbook_and_replaces_score_cells(tmp_path: Path) -> None:
    output = tmp_path / "filled.xlsx"
    render_scorecard_workbook(_manual_scores(), output)
    with ZipFile(output) as archive:
        assert "xl/styles.xml" in archive.namelist()
        assert "xl/sharedStrings.xml" in archive.namelist()
    assert _xlsx_numeric_cells(output)["D19"] == 72
