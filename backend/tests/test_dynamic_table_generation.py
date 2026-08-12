from pathlib import Path

from docx import Document

from docx_engine import DocxGenerator
from validator import Validator
from validator.mapper import InputDataMapper
from app.models.report import ReportTemplateType
from app.services.report_validator import normalize_dynamic_tables


def test_online_rows_beyond_template_capacity_are_generated_and_validated(tmp_path: Path) -> None:
    templates = Path(__file__).resolve().parents[1] / "app" / "templates"
    rows = [
        {"0": f"部门{i}", "1": i, "2": f"主要职能{i}", "3": f"负责人{i}"}
        for i in range(1, 8)
    ]
    data = {
        "cover_strategy_stock_discretionary": True,
        "__dynamic_tables": {"2": rows},
    }
    output = tmp_path / "dynamic-departments.docx"
    generation = DocxGenerator(
        templates / "private_fund_template.docx",
        templates / "private_fund_manifest.json",
    ).generate(data, output)

    document = Document(output)
    assert len(document.tables[1].rows) == 8  # one header plus seven online rows
    assert [cell.text for cell in document.tables[1].rows[-1].cells] == [
        "部门7", "7", "主要职能7", "负责人7",
    ]
    assert generation.summary.failed_fields == 0
    assert generation.summary.paragraph_validation.success

    validation = Validator(
        templates / "private_fund_template.docx", profile="private_2026"
    ).validate(output, data)
    assert validation.success


def test_cta_json_rows_keep_all_products_and_normalize_percentages() -> None:
    products = [
        {
            "name": f"CTA产品{i}", "established": f"202{i}.1.1",
            "annual_return": f"{i}.25%", "max_dd": f"{i}.5%", "aum": str(i),
        }
        for i in range(1, 6)
    ]
    mapped = InputDataMapper().map({
        "header": {"strategy_type": ["期货量化趋势"]},
        "section_1_9_strategy": {"managed_futures": {"representative_products": products}},
    }, "private_2026")

    rows = mapped["__dynamic_tables"]["9"]
    assert len(rows) == 5
    normalized = normalize_dynamic_tables(
        ReportTemplateType.PRIVATE_FUND, {"9": rows}
    )["9"]
    assert normalized[0]["2"] == "1.25"
    assert normalized[0]["3"] == "1.5"
    assert normalized[-1]["2"] == "5.25"


def test_short_dynamic_tables_drop_unused_template_rows_and_add_units(tmp_path: Path) -> None:
    templates = Path(__file__).resolve().parents[1] / "app" / "templates"
    data = {
        "cover_strategy_futures_quant_trend": True,
        "__dynamic_tables": {
            "2": [{"0": "投资部", "1": "8", "2": "投资研究", "3": "负责人"}],
            "9": [{"0": "CTA一号", "1": "2024.1.1", "2": "12.5", "3": "3.2", "4": "2.6"}],
        },
    }
    output = tmp_path / "trimmed.docx"
    generation = DocxGenerator(
        templates / "private_fund_template.docx",
        templates / "private_fund_manifest.json",
    ).generate(data, output)

    document = Document(output)
    assert len(document.tables[1].rows) == 2
    assert len(document.tables[8].rows) == 2
    assert document.tables[8].rows[1].cells[2].text == "12.5%"
    assert document.tables[8].rows[1].cells[4].text == "2.6亿"
    assert generation.summary.paragraph_validation.success

    validation = Validator(
        templates / "private_fund_template.docx", profile="private_2026"
    ).validate(output, data)
    assert validation.success


def test_all_strategy_branches_map_to_semantic_bookmarks() -> None:
    mapped = InputDataMapper().map({
        "header": {"report_date": "2026-06-26"},
        "section_1_1_basic_info_table": {
            "established_date": "2011-05-27",
            "other_whitelist": "招商银行、建设银行",
        },
        "section_1_9_strategy": {
            "quantitative": {"_applicable": True, "quant_methodology": "量化方法"},
            "managed_futures": {
                "_applicable": True,
                "position_sizing_method": "头寸方法",
                "data_sources_used": "数据源",
            },
            "fixed_income": {"_applicable": True, "investment_scope": "债券范围"},
            "options_arbitrage": {"_applicable": True, "strategy_logic_and_sub_weights": "期权逻辑"},
        },
    }, "private_2026")

    assert mapped["cover_report_date"] == "2026.6.26"
    assert mapped["table_1_row2_col1"].startswith("2011年5月27日")
    assert mapped["table_1_row21_col1"] == "是，招商银行、建设银行"
    assert mapped["strat_quant_methodology"] == "量化方法"
    assert mapped["strat_cta_position_sizing"] == "头寸方法"
    assert "数据源" in mapped["strat_cta_research_platform_data"]
    assert mapped["strat_bond_investment_scope"] == "债券范围"
    assert mapped["strat_option_strategy_logic_weights"] == "期权逻辑"


def test_other_product_strategy_marks_the_word_cover_checkbox(tmp_path: Path) -> None:
    templates = Path(__file__).resolve().parents[1] / "app" / "templates"
    output = tmp_path / "other-strategy.docx"
    DocxGenerator(
        templates / "private_fund_template.docx",
        templates / "private_fund_manifest.json",
    ).generate({"cover_strategy_other_text": "事件驱动"}, output)

    document = Document(output)
    cover = "\n".join(paragraph.text for paragraph in document.paragraphs[:20])
    assert "☑其他投资策略（事件驱动）" in cover
