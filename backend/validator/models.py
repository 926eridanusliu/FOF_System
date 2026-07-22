from __future__ import annotations

import json
from dataclasses import asdict, dataclass, field
from pathlib import Path
from typing import Any


@dataclass
class FieldResult:
    field: str
    expected: str
    actual: str
    status: str
    location: str
    message: str = ""

    def to_dict(self) -> dict[str, Any]:
        return asdict(self)


@dataclass
class FormatIssue:
    field: str
    location: str
    kind: str
    expected: str
    actual: str
    message: str

    def to_dict(self) -> dict[str, Any]:
        return asdict(self)


@dataclass
class TableIssue:
    table: int
    kind: str
    expected: str
    actual: str
    message: str

    def to_dict(self) -> dict[str, Any]:
        return asdict(self)


@dataclass
class Report:
    generated_docx: str
    profile: str
    matched: int
    missing: int
    mismatched: int
    extra: int
    empty_expected: int
    format_issue_count: int
    table_issue_count: int
    field_results: list[FieldResult] = field(default_factory=list)
    format_issues: list[FormatIssue] = field(default_factory=list)
    table_issues: list[TableIssue] = field(default_factory=list)
    notes: list[str] = field(default_factory=list)

    @property
    def success(self) -> bool:
        return not any([
            self.missing,
            self.mismatched,
            self.extra,
            self.format_issue_count,
            self.table_issue_count,
        ])

    def to_dict(self) -> dict[str, Any]:
        data = asdict(self)
        data["success"] = self.success
        return data

    def to_json(self, path: str | Path | None = None) -> str:
        content = json.dumps(self.to_dict(), ensure_ascii=False, indent=2)
        if path is not None:
            target = Path(path)
            target.parent.mkdir(parents=True, exist_ok=True)
            target.write_text(content, encoding="utf-8")
        return content

    def to_docx(self, path: str | Path) -> Path:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt

        target = Path(path)
        target.parent.mkdir(parents=True, exist_ok=True)
        document = Document()
        title = document.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("文档验证报告")
        run.bold = True
        run.font.size = Pt(18)
        document.add_paragraph(
            f"结果：{'通过' if self.success else '未通过'}；模板类型：{self.profile}；"
            f"一致：{self.matched}；遗漏：{self.missing}；不一致：{self.mismatched}；"
            f"多余：{self.extra}；格式异常：{self.format_issue_count}；"
            f"表格异常：{self.table_issue_count}。"
        )
        table = document.add_table(rows=1, cols=6)
        table.style = "Table Grid"
        for cell, text in zip(
            table.rows[0].cells,
            ["状态", "字段", "期望值", "实际值", "位置", "说明"],
        ):
            cell.text = text
        for item in self.field_results:
            row = table.add_row().cells
            for cell, text in zip(row, [
                item.status, item.field, item.expected, item.actual,
                item.location, item.message,
            ]):
                cell.text = text
        if self.format_issues:
            document.add_heading("格式异常", level=1)
            for issue in self.format_issues:
                document.add_paragraph(
                    f"{issue.field}｜{issue.kind}｜{issue.location}｜"
                    f"期望：{issue.expected}｜实际：{issue.actual}｜{issue.message}"
                )
        if self.table_issues:
            document.add_heading("表格异常", level=1)
            for issue in self.table_issues:
                document.add_paragraph(
                    f"表{issue.table}｜{issue.kind}｜期望：{issue.expected}｜"
                    f"实际：{issue.actual}｜{issue.message}"
                )
        if self.notes:
            document.add_heading("说明", level=1)
            for note in self.notes:
                document.add_paragraph(note)
        document.save(target)
        return target
